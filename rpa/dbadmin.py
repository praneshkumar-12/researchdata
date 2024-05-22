from rpa.models import Publications
from rpa.forms import PublicationsForm
from rpa.models import Users
from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.db.models import Count
import random
import string
import django.db.utils
import os
import json
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import rpa.extractor.extractor as Extractor


def admin_home(request):
    papers = Publications.objects.all()

    publication_list = []

    name = str(request.session.get("FACULTY_NAME"))

    if name is None or name != "admin" or name == str(None):
        return redirect("/rpa/login")

    for paper in papers:
        publication_list.append(paper)

    publication_list.sort(reverse=True, key=lambda x: x.end_academic_year)

    context = {"papers": publication_list, "name": name}

    return render(request, "admin_home.html", context)


def admin_dashboard(request):
    papers = Publications.objects.all()

    publication_list = []

    facs = []

    faculties = Users.objects.all().order_by("staff_name")

    for row in faculties:
        facs.append(row.staff_name)

    name = str(request.session.get("FACULTY_NAME"))

    if name is None or name != "admin" or name == "None":
        return redirect("/rpa/login")

    form = PublicationsForm()

    for paper in papers:
        first_author = paper.first_author if paper.first_author else ""
        second_author = paper.second_author if paper.second_author else ""
        third_author = paper.third_author if paper.third_author else ""
        other_authors = paper.other_authors if paper.other_authors else ""

        if not paper.second_author:
            paper.second_author = "NULL"
        if not paper.third_author:
            paper.third_author = "NULL"
        if not paper.other_authors:
            paper.other_authors = "NULL"
        if not paper.is_student_author:
            paper.is_student_author = "NULL"
        if not paper.student_name:
            paper.student_name = "NULL"
        if not paper.student_batch:
            paper.student_batch = "NULL"
        if not paper.specification:
            paper.specification = "NULL"
        if not paper.publication_type:
            paper.publication_type = "NULL"
        if not paper.publication_name:
            paper.publication_name = "NULL"
        if not paper.publisher:
            paper.publisher = "NULL"
        if not paper.year_of_publishing:
            paper.year_of_publishing = "NULL"
        if not paper.month_of_publishing:
            paper.month_of_publishing = "NULL"
        if not paper.page_number:
            paper.page_number = "NULL"
        if not paper.indexing:
            paper.indexing = "NULL"
        if not paper.quartile:
            paper.quartile = "NULL"
        if not paper.url:
            paper.url = "NULL"
        if not paper.issn:
            paper.issn = "NULL"
        if not paper.front_page_path:
            paper.front_page_path = "NULL"

        publication_list.append(paper)

    publication_list.sort(reverse=True, key=lambda x: x.end_academic_year)

    context = {
        "papers": publication_list,
        "form": form,
        "new_sno": f"{int(len(publication_list)) + 1}",
        "name": name,
        "faculties": facs,
    }

    # paper_records =
    return render(request, "admin_dashboard.html", context)


def admin_view_paper_details(request, paperid):
    if request.method != "GET":
        return render(request, "error.html")

    if not paperid:
        return render(request, "error.html")

    try:
        paper = Publications.objects.get(uniqueid=paperid)
    except Publications.DoesNotExist:
        return render(request, "error.html")

    name = str(request.session.get("FACULTY_NAME"))

    if name.lower().strip() != "admin":
        return render(
            request,
            "custom_error.html",
            {
                "error_title": "Unauthorized!",
                "error_message": "You are unauthorized to view the details of this publication.",
            },
        )

    return render(request, "admin_publication.html", {"paper": paper, "name": name})


def admin_upload_paper(request, uniqueid):
    if request.method == "GET":
        try:
            publ = Publications.objects.get(uniqueid=uniqueid)
        except Publications.DoesNotExist:
            return render(request, "error.html")

        name = str(request.session.get("FACULTY_NAME"))

        if name.lower() != "admin":
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Unauthorized!",
                    "error_message": "You are unauthorized to view the details of this publication.",
                },
            )

        return render(
            request, "admin_upload.html", {"title": publ.title, "uniqueid": uniqueid}
        )
    else:
        publ = Publications.objects.get(uniqueid=uniqueid)
        name = str(request.session.get("FACULTY_NAME"))

        if name.lower() != "admin":
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Unauthorized!",
                    "error_message": "You are unauthorized to view the details of this publication.",
                },
            )
        try:
            if not os.path.exists("rpa/static/upload/"):
                os.mkdir("rpa/static/upload/")

            upload_file = request.FILES["file"]

            complete_path = "rpa/static/upload/" + uniqueid.strip() + ".pdf"

            publ = Publications.objects.get(uniqueid=uniqueid)

            with open(complete_path, "wb+") as destination:
                for chunk in upload_file.chunks():
                    destination.write(chunk)

            publ.front_page_path = complete_path

            publ.save()

            return render(
                request,
                "upload.html",
                {"alertmessage": "Upload successful!", "uniqueid": uniqueid},
            )
        except Exception as e:
            return render(
                request,
                "upload.html",
                {"alertmessage": str(e), "reload": "yes", "uniqueid": uniqueid},
            )


def admin_remove_upload(request):
    uniqueid = request.POST.get("uniqueid")

    name = str(request.session.get("FACULTY_NAME", ""))

    if name != "admin":
        return HttpResponse("Unauthorized")

    if not Publications.objects.filter(uniqueid=uniqueid):
        return HttpResponse("Paper not found!")

    publ = Publications.objects.get(uniqueid=uniqueid)

    complete_path = "rpa/static/upload/" + uniqueid.strip() + ".pdf"

    if os.path.exists(complete_path):
        os.remove(complete_path)
    else:
        return HttpResponse("File not found!")

    publ.front_page_path = None

    publ.save()

    return HttpResponse("OK")


def admin_verify_paper(request):
    uniqueid = request.POST.get("uniqueid")

    name = str(request.session.get("FACULTY_NAME"))

    if name != "admin":
        return HttpResponse("Unauthorized")

    if not Publications.objects.filter(uniqueid=uniqueid):
        return HttpResponse("Paper not found!")

    current_paper = Publications.objects.get(uniqueid=uniqueid)
    current_paper.admin_verified = "True"
    current_paper.save()

    return HttpResponse("OK")


def admin_update_paper(request):
    uniqueid = request.POST.get("uniqueid")

    name = str(request.session.get("FACULTY_NAME"))

    if name != "admin":
        return HttpResponse("Unauthorized")

    if not Publications.objects.filter(uniqueid=uniqueid):
        return HttpResponse("Paper not found!")

    updateData = dict(request.POST)

    updates = {}

    for key, value in updateData.items():
        if (
            key != "csrfmiddlewaretoken"
            and key != "uniqueid"
            and key != "academic_year"
        ):
            if key == "student_author":
                key = "is_student_author"
            if key == "publication_year":
                key = "year_of_publishing"
            if key == "publication_month":
                key = "month_of_publishing"
            if key == "citations":
                key = "citation"
            if key == "page_numbers":
                key = "page_number"
            try:
                updates[key] = value[0]
            except IndexError:
                updates[key] = "NULL"

    academic_year = request.POST.get("academic_year").replace(" -", "").split(" ")

    start_academic_month = academic_year[0]
    start_academic_year = academic_year[1]
    end_academic_month = academic_year[2]
    end_academic_year = academic_year[3]

    updates["start_academic_year"] = int(start_academic_year)
    updates["start_academic_month"] = start_academic_month
    updates["end_academic_year"] = int(end_academic_year)
    updates["end_academic_month"] = end_academic_month
    updates["volume"] = int(updates["volume"])
    updates["citation"] = int(updates["citation"])
    updates["year_of_publishing"] = int(updates["year_of_publishing"])

    updates["verified"] = "False"
    updates["admin_verified"] = "False"

    Publications.objects.filter(uniqueid=uniqueid).update(**updates)

    return HttpResponse("OK")


def admin_verification(request):
    papers = Publications.objects.all()

    publication_list = []

    name = str(request.session.get("FACULTY_NAME"))

    if name is None or name != "admin" or name == str(None):
        return redirect("/rpa/login")

    for paper in papers:
        publication_list.append(paper)

    publication_list.sort(reverse=True, key=lambda x: x.end_academic_year)

    context = {"papers": publication_list, "name": name}

    return render(request, "admin_verification.html", context)


def admin_get_doi(request):
    if request.method == "POST":
        doi = request.POST.get("doi")
        ay = request.POST.get("AY")

        title = Extractor.get_title(doi)

        if not title:
            return render(request, "admin_get_title.html", {"doi": doi, "AY": ay})

        search_query = title

        try:
            result = Extractor.main(search_query, ay)
        except (IndexError, KeyError):
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Uh-Oh! Unable to fetch!",
                    "error_message": "Cannot fetch data for the given term!",
                },
            )

        if (result is None) or (result == False):
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Uh-Oh! Unable to fetch!",
                    "error_message": "Cannot fetch data for the given term!",
                },
            )
        else:
            if result.get("volume") is None:
                result["volume"] = 0
            if result.get("citation") is None:
                result["citation"] = 0
            
            return render(request, "admin_add_publication.html", {"result": result})

    name = request.session.get("FACULTY_NAME", "")

    if name != "admin":
        return redirect("/rpa/user/error")

    return render(request, "admin_get_doi.html", {"name": name})


def admin_get_title(request):
    if request.method == "POST":
        title = request.POST.get("title", "")
        doi = request.POST.get("doi", "")
        ay = request.POST.get("AY", "")

        search_query = title + doi

        try:
            result = Extractor.main(search_query, ay)
        except (IndexError, KeyError):
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Uh-Oh! Unable to fetch!",
                    "error_message": "Cannot fetch data for the given term!",
                },
            )

        if (result is None) or (result == False):
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Uh-Oh! Unable to fetch!",
                    "error_message": "Cannot fetch data for the given term!",
                },
            )
        else:
            if result.get("volume") is None:
                result["volume"] = 0
            if result.get("citation") is None:
                result["citation"] = 0
            
            return render(request, "admin_add_publication.html", {"result": result})

    name = request.session.get("FACULTY_NAME", "")

    if name != "admin":
        return redirect("/rpa/user/error")

    return render(request, "admin_get_title.html", {"name": name})


def admin_insert_paper(request):
    data_to_be_inserted = {}

    faculty_name = request.session["FACULTY_NAME"]

    if faculty_name != "admin":
        return HttpResponse("Unauthorized")

    uniqueid = request.POST.get("uniqueid")

    exists = Publications.objects.filter(uniqueid=uniqueid)

    if exists:
        return HttpResponse("Paper already exists! Cannot add paper.")

    data_to_be_inserted = {}

    for key, value in request.POST.items():
        if (
            key != "csrfmiddlewaretoken"
            and key != "uniqueid"
            and key != "academic_year"
        ):
            if key == "student_author":
                key = "is_student_author"
            if key == "publication_year":
                key = "year_of_publishing"
            if key == "publication_month":
                key = "month_of_publishing"
            if key == "citations":
                key = "citation"
            if key == "page_numbers":
                key = "page_number"
            try:
                data_to_be_inserted[key] = value
            except IndexError:
                data_to_be_inserted[key] = "NULL"

    data_to_be_inserted["uniqueid"] = uniqueid

    academic_year = request.POST.get("academic_year").replace(" -", "").split(" ")

    start_academic_month = academic_year[0]
    start_academic_year = academic_year[1]
    end_academic_month = academic_year[2]
    end_academic_year = academic_year[3]

    data_to_be_inserted["start_academic_year"] = int(start_academic_year)
    data_to_be_inserted["start_academic_month"] = start_academic_month
    data_to_be_inserted["end_academic_year"] = int(end_academic_year)
    data_to_be_inserted["end_academic_month"] = end_academic_month
    data_to_be_inserted["volume"] = int(data_to_be_inserted["volume"])
    data_to_be_inserted["citation"] = int(data_to_be_inserted["citation"])
    data_to_be_inserted["year_of_publishing"] = int(
        data_to_be_inserted["year_of_publishing"]
    )

    new_record = Publications(**data_to_be_inserted)
    try:
        new_record.save()
    except django.db.utils.IntegrityError:
        return HttpResponse("Paper already exists! Cannot add paper.")

    return HttpResponse("OK")


def admin_manually_insert_paper(request):
    if request.method == "GET":
        name = request.session.get("FACULTY_NAME", "")

        if name != "admin":
            return render(
                request,
                "custom_error.html",
                {
                    "error_title": "Unauthorized!",
                    "error_message": "You are unauthorized to do this operation!",
                },
            )

        return render(request, "admin_manually_add_paper.html", {"name": name})

    elif request.method == "POST":
        name = str(request.session.get("FACULTY_NAME"))

        if name != "admin":
            return HttpResponse("Unauthorized")

        updateData = dict(request.POST)

        updates = {}

        for key, value in updateData.items():
            if (
                key != "csrfmiddlewaretoken"
                and key != "uniqueid"
                and key != "academic_year"
            ):
                if key == "student_author":
                    key = "is_student_author"
                if key == "publication_year":
                    key = "year_of_publishing"
                if key == "publication_month":
                    key = "month_of_publishing"
                if key == "citations":
                    key = "citation"
                if key == "page_numbers":
                    key = "page_number"
                try:
                    updates[key] = value[0]
                except IndexError:
                    updates[key] = "NULL"

        academic_year = request.POST.get("academic_year").replace(" -", "").split(" ")

        start_academic_month = academic_year[0]
        start_academic_year = academic_year[1]
        end_academic_month = academic_year[2]
        end_academic_year = academic_year[3]

        updates["start_academic_year"] = int(start_academic_year)
        updates["start_academic_month"] = start_academic_month
        updates["end_academic_year"] = int(end_academic_year)
        updates["end_academic_month"] = end_academic_month
        updates["volume"] = int(updates["volume"])
        updates["citation"] = int(updates["citation"])
        updates["year_of_publishing"] = int(updates["year_of_publishing"])

        updates["verified"] = "False"
        updates["admin_verified"] = "False"

        uniqueid = (
            str(start_academic_year)
            + str(start_academic_month)
            + "".join(random.choices(string.ascii_letters, k=7))
        )

        updates["uniqueid"] = uniqueid

        if Publications.objects.filter(uniqueid=uniqueid):
            return HttpResponse("Paper already exists! Cannot add paper.")

        new_record = Publications(**updates)

        try:
            new_record.save()
        except django.db.utils.IntegrityError:
            return HttpResponse("Paper already exists! Cannot add paper.")

        return HttpResponse("OK")


def admin_delete_paper(request):
    uniqueid = request.POST.get("uniqueid")


    if not Publications.objects.filter(uniqueid=uniqueid):
        return HttpResponse("Paper not found!")

    publication = Publications.objects.get(uniqueid=uniqueid)

    publication.delete()

    return HttpResponse("OK")

def convert_to_dict(row):
    row_data = {
                    "uniqueid" : row[1] if row[1] and row[1].lower() != 'null' and row[1].lower() != 'none' else "",
                    "title": row[2] if row[2] and row[2].lower() != 'null' and row[2].lower() != 'none' else "",
                    "AY": row[3] if row[3] and row[3].lower() != 'null' and row[3].lower() != 'none' else "",
                    "first_author" : row[4] if row[4] and row[4].lower() != 'null' and row[4].lower() != 'none' else "",
                    "second_author" : row[5] if row[5] and row[5].lower() != 'null' and row[5].lower() != 'none' else "",
                    "third_author" : row[6] if row[6] and row[6].lower() != 'null' and row[6].lower() != 'none' else "",
                    "other_authors" : row[7] if row[7] and row[7].lower() != 'null' and row[7].lower() != 'none' else "",
                    # "is_student_author" : row['is_student_author'],
                    # "student_name" : row['student_name'],
                    # "student_batch" : row['student_batch'],
                    "specification" : row[11] if row[11] and row[11].lower() != 'null' and row[11].lower() != 'none' else "",
                    "publication_type" : row[12] if row[12] and row[12].lower() != 'null' and row[12].lower() != 'none' else "",
                    "publication_name" : row[13] if row[13] and row[13].lower() != 'null' and row[13].lower() != 'none' else "",
                    "publisher" : row[14] if row[14] and row[14].lower() != 'null' and row[14].lower() != 'none' else "",
                    "year_of_publishing" : row[15] if row[15] and row[15].lower() != 'null' and row[15].lower() != 'none' else "",
                    "month_of_publishing" : row[16] if row[16] and row[16].lower() != 'null' and row[16].lower() != 'none' else "",
                    "volume" : row[17] if row[17] and row[17].lower() != '0' else "",
                    "page_number" : row[18] if row[18] and row[18].lower() != 'null' and row[18].lower() != 'none' else "",
                    "indexing" : row[19] if row[19] and row[19].lower() != 'null' and row[19].lower() != 'none' else "",
                    "quartile" : row[20] if row[20] and row[20].lower() != 'null' and row[20].lower() != 'none' else "",
                    # "citation" : row['citation'],
                    "doi" : row[22] if row[22] and row[22].lower() != 'null' and row[22].lower() != 'none' else "",
                    # "front_page_path" : row['front_page_path'],
                    "url" : row[24] if row[24] and row[24].lower() != 'null' and row[24].lower() != 'none' else "",
                    "issn" : row[25] if row[25] and row[25].lower() != 'null' and row[25].lower() != 'none' else "",
                    # "verified" : row['verified'],
                    # "admin_verified" : row['admin_verified'],
                }
    return row_data

def IEEEFormat(paper):
    '''
        Author initials. Last name, “Article title,” Journal Name, vol. Volume, no. Number, pp. Page range, Month Year, DOI.
    '''

    months = {
        1: "Jan",
        2: "Feb",
        3: "Mar",
        4: "Apr",
        5: "May",
        6: "Jun",
        7: "Jul",
        8: "Aug",
        9: "Sep",
        10: "Oct",
        11: "Nov",
        12: "Dec",
    }


    format_string = ''

    if paper.get("first_author"):
        format_string += str(paper.get("first_author")) + str(", ")
    
    if paper.get("second_author"):
        format_string += str(paper.get("second_author")) + str(", ")
    
    if paper.get("third_author"):
        format_string += str(paper.get("third_author")) + str(", ")

    if paper.get("other_authors"):
        format_string += str(paper.get("other_authors")) + str(", ")

    if paper.get("title"):
        format_string += str("\"") + str(paper.get("title")) + str("\", ")
    
    if paper.get("publication_name"):
        format_string += str(paper.get("publication_name")) + str(", ")
    
    if paper.get("publisher"):
        format_string += str(paper.get("publisher")) + str(", ")
    
    if paper.get("volume") and paper.get("volume") != "0":
        format_string += str("vol. ") + str(paper.get("volume")) + str(", ")
    
    if paper.get("page_number"):
        format_string += str("pp. ") + str(paper.get("page_number")) + str(", ")
    
    if paper.get("year_of_publishing"):
        if paper.get("month_of_publishing"):
            format_string += str(months.get(int(paper.get("month_of_publishing")))) + str(" ") + str(paper.get("year_of_publishing")) + str(", ")
        else:
            format_string += str(paper.get("year_of_publishing")) + str(", ")
    
    if paper.get("doi"):
        format_string += str("DOI: ") + str(paper.get("doi")) + str(", ")
    
    if paper.get("issn"):
        format_string += str("ISSN: ") + str(paper.get("issn")) + str(", ")
    
    if paper.get("indexing") or paper.get("quartile"):
        if paper.get("indexing") and paper.get("quartile"):
            quartile = paper.get("quartile")
            if paper.get("quartile").lower() == "q4":
                quartile = "< Q3"
            format_string += str("(Indexed in: ") + str(paper.get("indexing")) + str(" Quartile: ") + str(quartile) + str("), ")
        elif paper.get("indexing"):
            format_string += str("(Indexed in: ") + str(paper.get("indexing")) + str("), ") 
        elif paper.get("quartile"):
            quartile = paper.get("quartile")
            if paper.get("quartile").lower() == "q4":
                quartile = "< Q3"
            format_string += str("(Quartile: ") + str(quartile) + str("), ")
    
    if paper.get("url"):
        format_string += str("URL: ") + str(paper.get("url"))

    format_string = format_string.rstrip(", ")
    
    return ' '.join(format_string.split())

    

def generate_word_document(data):
    sorted_data = sorted(data, key=lambda x: x['AY'])

    journals = {}
    conference = {}
    book_series = {}


    for data in sorted_data:
        if data["publication_type"].lower().strip() == "journal":
            if data["AY"] not in journals:
                journals[data["AY"]] = {}

                print(f'{data["quartile"]=}, {bool(data["quartile"])}')

                if data["quartile"] and data["quartile"].lower() != "null" and data["quartile"].lower() != "none":
                    if data["quartile"].lower() == "q4":
                        quartile = "< Q3"
                    else:
                        quartile = data["quartile"]
                else:
                    quartile = "< Q3"

                if quartile not in journals[data["AY"]]:
                    journals[data["AY"]][quartile] = [data]
                else:
                    journals[data["AY"]][quartile].append(data)
            else:
                print(f'{data["quartile"]=}, {bool(data["quartile"])}')

                if data["quartile"] and data["quartile"].lower() != "null" and data["quartile"].lower() != "none":
                    if data["quartile"].lower() == "q4":
                        quartile = "< Q3"
                    else:
                        quartile = data["quartile"]
                else:
                    quartile = "< Q3"
                
                if quartile not in journals[data["AY"]]:
                    journals[data["AY"]][quartile] = [data]
                else:
                    journals[data["AY"]][quartile].append(data)

        elif data["publication_type"].lower().strip() == "conference":
            if data["AY"] not in conference:
                conference[data["AY"]] = [data]
            else:
                conference[data["AY"]].append(data)
        elif data["publication_type"].lower().strip() == "book chapter":
            if data["AY"] not in book_series:
                book_series[data["AY"]] = [data]
            else:
                book_series[data["AY"]].append(data)
            
    # Create a new Word document
    doc = docx.Document()

    heading1 = doc.add_heading('Department of Information Technology', level=1)
    heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading1.runs[0]
    run.font.size = Pt(24)
    run.bold = True
    run.underline = True

    heading1 = doc.add_heading('Research Publications', level=1)
    heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading1.runs[0]
    run.font.size = Pt(20)
    run.bold = True
    run.underline = True

    if journals:
        journal_heading = doc.add_heading('Journals', level=1)
        run = journal_heading.runs[0]
        run.underline = True

        for AY, details in journals.items():
            if details:
                AY_heading = doc.add_heading(AY, level=2) 
                run = AY_heading.runs[0]
                run.underline = True

                details = sort_quartile_records(details)

                for quartile, papers in details.items():
                    if papers:
                        quartile_heading = doc.add_heading(f"Publications under {quartile}", level=3)
                        run = quartile_heading.runs[0]
                        run.underline = True

                        for idx, paper in enumerate(papers):
                            paragraph = doc.add_paragraph(str(idx+1) + ". " + IEEEFormat(paper))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if book_series:
        book_series_heading = doc.add_heading('Book Series', level=1)
        run = book_series_heading.runs[0]
        run.underline = True

        for AY, details in book_series.items():
            if details:
                AY_heading = doc.add_heading(AY, level=2)
                run = AY_heading.runs[0]
                run.underline = True
            
                for idx, paper in enumerate(details):
                    paragraph = doc.add_paragraph(str(idx+1) + ". " + IEEEFormat(paper))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if conference:
        conference_heading = doc.add_heading('Conference', level=1)
        run = conference_heading.runs[0]
        run.underline = True

        for AY, details in conference.items():
            if details:
                AY_heading = doc.add_heading(AY, level=2)
                run = AY_heading.runs[0]
                run.underline = True
            
                for idx, paper in enumerate(details):
                    paragraph = doc.add_paragraph(str(idx+1) + ". " + IEEEFormat(paper))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.save("research_publications.docx")

def admin_get_word(request):
    if request.method == "POST":
        table_data = request.POST.get("data")
        table_dict = json.loads(f'{{"rows": {table_data}}}')

        table_rows = table_dict.get("rows")

        all_row_data = []

        for row in table_rows:
            row_data = convert_to_dict(row)
            all_row_data.append(row_data)

        generate_word_document(all_row_data)

    return JsonResponse({})

def sort_quartile_records(input_dict):
    # Define the desired order of keys
    desired_order = ['Q1', 'Q2', 'Q3', '< Q3']
    
    # Create a new dictionary with keys sorted according to the desired order
    sorted_dict = {key: input_dict[key] for key in desired_order if key in input_dict}
    
    return sorted_dict


def admin_get_charts(request):
    if request.method == "GET":
        name = str(request.session.get("FACULTY_NAME"))

        if name is None or name != "admin" or name == str(None):
            return redirect("/rpa/login")
            
        # Fetch data for donut chart
        chart_records = {}
        all_records = Publications.objects.all()
        for record in all_records:
            if (
                not record.indexing
                or record.indexing == "NULL"
                or record.indexing == "None"
            ):
                key = "Others"
            else:
                indices = record.indexing.split(", ")
                for index in indices:
                    if index not in ['Scopus', 'Web of Sciences']:
                        key = "Others"
                        chart_records[key] = chart_records.get(key, 0) + 1
                        continue
                    key = index
                    chart_records[key] = chart_records.get(key, 0) + 1

        donut_labels = list(chart_records.keys())
        donut_values = list(chart_records.values())

        # Fetch data for bar chart based on academic year/month
        bar_data = Publications.objects.values(
            "start_academic_year",
            "start_academic_month",
            "end_academic_year",
            "end_academic_month",
        ).annotate(total=Count("uniqueid")).order_by("start_academic_year")
        bar_labels = []
        bar_values = []
        for data in bar_data:
            start_academic_year = data["start_academic_year"]
            start_academic_month = data["start_academic_month"]
            end_academic_year = data["end_academic_year"]
            end_academic_month = data["end_academic_month"]
            label = f"{start_academic_month} {start_academic_year} - {end_academic_month} {end_academic_year}"
            bar_labels.append(label)
            bar_values.append(data["total"])

        # Fetch data for another bar chart based on publication types
        publication_types_data = Publications.objects.values("publication_type").annotate(
            total=Count("uniqueid")
        )
        publication_type_labels = []
        publication_type_values = []
        for pub_type_data in publication_types_data:
            publication_type_labels.append(pub_type_data["publication_type"])
            publication_type_values.append(pub_type_data["total"])

        temp = {}

        for i in range(len(publication_type_labels)):
            temp[publication_type_labels[i]] = publication_type_values[i]
        
        publication_type_values = [temp]
        publication_type_labels = ["Overall"]

        # Fetch data for quartile bar chart
        quartile_records = {}
        quartile_labels = []
        quartile_values = []
        qr_records = Publications.objects.all()
        for record in qr_records:
            if (
                not record.quartile
                or record.quartile == "NULL"
                or record.quartile == "None"
                or record.quartile == "Q4"
            ):
                key = "< Q3"
            else:
                key = record.quartile
            quartile_records[key] = quartile_records.get(key, 0) + 1

        quartile_records = sort_quartile_records(quartile_records)


        quartile_labels = list(quartile_records.keys())
        quartile_values = list(quartile_records.values())


        yearly_quartiles = {}
        for data in bar_data:
            year = f"{data['start_academic_year']} - {data['end_academic_year']}"
            if year not in yearly_quartiles:
                yearly_quartiles[year] = {'Q1': 0, 'Q2': 0, 'Q3': 0, "< Q3": 0}
            year_records = all_records.filter(
                start_academic_year=data["start_academic_year"],
                end_academic_year=data["end_academic_year"],
                start_academic_month=data["start_academic_month"],
                end_academic_month=data["end_academic_month"],
            )
            for record in year_records:
                if record.quartile and  record.quartile != 'NULL' and record.quartile in ("Q1", "Q2", "Q3"):
                    yearly_quartiles[year][record.quartile] += 1
                else:
                    yearly_quartiles[year]["< Q3"] += 1
        yearwise_label = list(yearly_quartiles.keys())
        yearwise_values = list(yearly_quartiles.values())
        
        # Prepare data for rendering
        data = {
            "donut_labels": donut_labels,
            "donut_values": donut_values,
            "bar_labels": bar_labels,
            "bar_values": bar_values,
            "publication_type_labels": publication_type_labels,
            "publication_type_values": publication_type_values,
            "quartile_labels": quartile_labels,
            "quartile_values": quartile_values,
            "name": name,
            "yearwise_label": yearwise_label,
            "yearwise_values": yearwise_values
        }

        return render(request, "charts.html", data)
    elif request.method == "POST":
        name = str(request.session.get("FACULTY_NAME"))

        if name is None or name != "admin" or name == str(None):
            return redirect("/rpa/login")
                
        AY = request.POST.get("AY")

        if AY == "all":
            return redirect('/rpa/dbadmin/charts')
        
        start, end = AY.split(" - ")
        start_academic_month, start_academic_year = start.split(" ")
        end_academic_month, end_academic_year = end.split(" ")

        # Fetch data for donut chart
        chart_records = {}
        all_records = Publications.objects.filter(start_academic_year=start_academic_year, end_academic_year=end_academic_year, start_academic_month=start_academic_month, end_academic_month=end_academic_month)
        for record in all_records:
            if (
                not record.indexing
                or record.indexing == "NULL"
                or record.indexing == "None"
            ):
                key = "Others"
            else:
                indices = record.indexing.split(", ")
                for index in indices:
                    if index not in ['Scopus', 'Web of Sciences']:
                        key = "Others"
                        chart_records[key] = chart_records.get(key, 0) + 1
                        continue
                    key = index
                    chart_records[key] = chart_records.get(key, 0) + 1

        donut_labels = list(chart_records.keys())
        donut_values = list(chart_records.values())

        # Fetch data for bar chart based on academic year/month
        bar_data = Publications.objects.values(
            "start_academic_year",
            "start_academic_month",
            "end_academic_year",
            "end_academic_month",
        ).annotate(total=Count("uniqueid")).order_by("start_academic_year")
        bar_labels = []
        bar_values = []
        for data in bar_data:
            start_academic_year = data["start_academic_year"]
            start_academic_month = data["start_academic_month"]
            end_academic_year = data["end_academic_year"]
            end_academic_month = data["end_academic_month"]
            label = f"{start_academic_month} {start_academic_year} - {end_academic_month} {end_academic_year}"
            bar_labels.append(label)
            bar_values.append(data["total"])

        # Fetch data for another bar chart based on publication types
        publication_types_data = all_records
        publication_type_labels = []
        publication_type_values = []
        count_data = {}
        for pub_type_data in publication_types_data:
            if pub_type_data.publication_type in count_data:
                count_data[pub_type_data.publication_type] += 1
            else:
                count_data[pub_type_data.publication_type] = 1

        publication_type_labels = list(count_data.keys())
        publication_type_values = list(count_data.values())

        temp = {}

        for i in range(len(publication_type_labels)):
            temp[publication_type_labels[i]] = publication_type_values[i]
        
        publication_type_values = [temp]
        publication_type_labels = [AY]

        # Fetch data for quartile bar chart
        quartile_records = {}
        quartile_labels = []
        quartile_values = []
        qr_records = all_records

        for record in qr_records:
            if (
                not record.quartile
                or record.quartile == "NULL"
                or record.quartile == "None"
                or record.quartile == "Q4"
            ):
                key = "< Q3"
            else:
                key = record.quartile
            quartile_records[key] = quartile_records.get(key, 0) + 1
        
        quartile_records = sort_quartile_records(quartile_records)

        quartile_labels = list(quartile_records.keys())
        quartile_values = list(quartile_records.values())

        yearly_quartiles = {}
        yearly_quartiles[AY] = {'Q1': 0, 'Q2': 0, 'Q3': 0, "< Q3": 0}
        for data in bar_data:
            # year = f"{data['start_academic_year']} - {data['end_academic_year']}"
            year_records = all_records.filter(
                start_academic_year=data["start_academic_year"],
                end_academic_year=data["end_academic_year"],
                start_academic_month=data["start_academic_month"],
                end_academic_month=data["end_academic_month"],
            )
            for record in year_records:
                if record.quartile and  record.quartile != 'NULL' and record.quartile in ("Q1", "Q2", "Q3"):
                    yearly_quartiles[AY][record.quartile] += 1
                else:
                    yearly_quartiles[AY]["< Q3"] += 1
        yearwise_label = list(yearly_quartiles.keys())
        yearwise_values = list(yearly_quartiles.values())
        
        # Prepare data for rendering
        data = {
            "donut_labels": donut_labels,
            "donut_values": donut_values,
            "bar_labels": bar_labels,
            "bar_values": bar_values,
            "publication_type_labels": publication_type_labels,
            "publication_type_values": publication_type_values,
            "quartile_labels": quartile_labels,
            "quartile_values": quartile_values,
            "name": name,
            "AY": AY,
            "yearwise_label": yearwise_label,
            "yearwise_values": yearwise_values,
        }

        return render(request, "charts.html", data)

